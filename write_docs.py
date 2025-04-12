import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
#from docx2pdf import convert


def write_docx(data: dict, name, idea):

    doc = Document()

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_logo = p_logo.add_run()

    run_logo.add_picture("outputs/logo.png", width=Inches(2.5))

    doc.add_paragraph("\n \n \n")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = p_title.add_run(f"{name} Documentation")
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.name = "Helvetica"

    p_powered = doc.add_paragraph()
    p_powered.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_powered = p_powered.add_run("Powered by Venture Force")
    run_powered.font.size = Pt(12)
    run_powered.font.name = "Helvetica"

    doc.add_page_break()

    title = doc.add_heading(name, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_picture("outputs/logo.png", width=Pt(2.5))

    doc.add_heading("User Project Idea", level=2)
    doc.add_paragraph(idea)

    print("Writing document...")
    print(data)

    for agent, responses in data.items():

        doc.add_heading(agent.upper(), level=1)

        for heading, content in responses.items():

            doc.add_heading(heading, level=2)
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(f"• {item}", style="List Bullet")
            elif isinstance(content, dict):
                for key, value in content.items():
                    doc.add_heading(key, level=3)
                    if isinstance(value, list):
                        for sub_item in value:
                            doc.add_paragraph(f"• {sub_item}",
                                              style="List Bullet")
                    else:
                        doc.add_paragraph(value)
            else:
                doc.add_paragraph(content)

    docx_filename = r"outputs/Project_Documentation.docx"
    doc.save(docx_filename)

    pdf_filename = "Project_Documentation.pdf"
    #convert(docx_filename)

    print(f"Document successfully saved as {pdf_filename}.")
